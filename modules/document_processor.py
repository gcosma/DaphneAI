# modules/document_processor.py - FIXED VERSION
# Improved PDF extraction with better cleaning and table handling

from datetime import datetime
from typing import List, Dict, Any, Optional
import logging
import re

# PDF processing with fallbacks
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logging.warning("PyPDF2 not available")

# DOCX processing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")


def clean_extracted_text(text: str) -> str:
    """
    ENHANCED: Clean text extracted from documents
    Removes PDF artifacts, duplicate text, and formatting issues
    """
    if not text:
        return ""
    
    # Fix UTF-8 encoding corruption
    text = text.replace('â€™', "'")
    text = text.replace('â€œ', '"')
    text = text.replace('â€', '"')
    text = text.replace('â€¢', '•')
    text = text.replace('â€"', '-')
    text = text.replace('â€¦', '...')
    
    # Fix common PDF extraction errors
    text = text.replace('ﬁ', 'fi').replace('ﬂ', 'fl')
    text = text.replace('ﬀ', 'ff').replace('ﬃ', 'ffi').replace('ﬄ', 'ffl')
    
    # Replace smart quotes
    text = text.replace('"', '"').replace('"', '"')
    text = text.replace(''', "'").replace(''', "'")
    text = text.replace('–', '-').replace('—', '-')
    text = text.replace('…', '...')
    
    # Remove unicode spaces
    text = text.replace('\u00A0', ' ')
    text = text.replace('\u2009', ' ')
    text = text.replace('\u200B', '')
    text = text.replace('\u202F', ' ')
    text = text.replace('\u2028', '\n')
    text = text.replace('\u2029', '\n\n')
    
    # ===================================================================
    # NEW: Remove PDF artifacts and repeated text
    # ===================================================================
    
    # Remove page numbers (various formats)
    text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
    text = re.sub(r'\n\s*Page\s+\d+\s*\n', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'\n\s*\d+\s+of\s+\d+\s*\n', '\n', text, flags=re.IGNORECASE)
    
    # Remove headers/footers that repeat on every page
    # Look for lines that appear multiple times
    lines = text.split('\n')
    line_counts = {}
    for line in lines:
        stripped = line.strip()
        if len(stripped) > 10 and len(stripped) < 100:
            line_counts[stripped] = line_counts.get(stripped, 0) + 1
    
    # Remove lines that appear more than 3 times (likely headers/footers)
    frequent_lines = {line for line, count in line_counts.items() if count > 3}
    filtered_lines = []
    for line in lines:
        if line.strip() not in frequent_lines:
            filtered_lines.append(line)
    
    text = '\n'.join(filtered_lines)
    
    # Remove "Recommendation N" followed immediately by "N" (duplicate numbering)
    text = re.sub(r'(Recommendation\s+\d+[a-z]?)\s+\1', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'(Recommendation\s+)(\d+[a-z]?)\s+\2', r'\1\2', text, flags=re.IGNORECASE)
    
    # Remove standalone numbers that are likely reference numbers
    text = re.sub(r'\n\s*\d+[a-z]?\s*\n', '\n', text)
    
    # Fix broken sentences (remove mid-sentence line breaks)
    text = re.sub(r'([a-z,;])\n([a-z])', r'\1 \2', text)
    
    # Clean up whitespace
    text = re.sub(r' {2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove lines that are just numbers or letters (artifacts)
    lines = text.split('\n')
    clean_lines = []
    for line in lines:
        stripped = line.strip()
        # Skip lines that are:
        # - Just numbers
        # - Just letters (a, b, c, etc.)
        # - Very short (< 10 chars) and don't look like real content
        if len(stripped) < 10:
            if stripped.isdigit() or (len(stripped) <= 3 and stripped.isalpha()):
                continue
        clean_lines.append(line)
    
    text = '\n'.join(clean_lines)
    
    return text.strip()


def process_uploaded_files(uploaded_files: List) -> List[Dict[str, Any]]:
    """Process uploaded files for document search"""
    if not uploaded_files:
        return []
    
    processed_docs = []
    
    for uploaded_file in uploaded_files:
        try:
            filename = uploaded_file.name
            file_size = len(uploaded_file.getvalue())
            
            if filename.lower().endswith('.pdf'):
                text = extract_pdf_text(uploaded_file)
            elif filename.lower().endswith('.txt'):
                text = extract_txt_text(uploaded_file)
            elif filename.lower().endswith('.docx'):
                text = extract_docx_text(uploaded_file)
            else:
                processed_docs.append({
                    'filename': filename,
                    'error': f'Unsupported file type',
                    'file_type': filename.split('.')[-1].lower() if '.' in filename else 'unknown',
                    'processed_at': datetime.now().isoformat()
                })
                continue
            
            # ENHANCED: Better text cleaning
            text = clean_extracted_text(text)
            
            doc = {
                'filename': filename,
                'text': text,
                'word_count': len(text.split()) if text else 0,
                'file_type': filename.split('.')[-1].lower() if '.' in filename else 'unknown',
                'file_size_mb': round(file_size / 1024 / 1024, 2),
                'processed_at': datetime.now().isoformat(),
                'processing_status': 'success'
            }
            
            processed_docs.append(doc)
            logging.info(f"Processed {filename}: {doc['word_count']} words")
            
        except Exception as e:
            logging.error(f"Error processing {filename}: {e}")
            processed_docs.append({
                'filename': filename,
                'error': str(e),
                'file_type': filename.split('.')[-1].lower() if '.' in filename else 'unknown',
                'processing_status': 'error',
                'processed_at': datetime.now().isoformat()
            })
    
    return processed_docs


def extract_table_as_structured_text(table: List[List[str]]) -> str:
    """Convert a table to structured text, preserving row relationships."""
    if not table or len(table) < 2:
        return ""
    
    # Clean table cells
    cleaned_table = []
    for row in table:
        cleaned_row = []
        for cell in row:
            if cell is None:
                cleaned_row.append("")
            else:
                cell_text = str(cell).strip()
                cell_text = re.sub(r'\s+', ' ', cell_text)
                cleaned_row.append(cell_text)
        cleaned_table.append(cleaned_row)
    
    headers = cleaned_table[0]
    header_text = ' '.join(headers).lower()
    
    # Check if recommendation table
    is_recommendation_table = any(kw in header_text for kw in [
        'recommendation', 'action', 'owner', 'timescale', 'improvement'
    ])
    
    if is_recommendation_table:
        return extract_recommendation_table(cleaned_table, headers)
    else:
        return extract_generic_table(cleaned_table, headers)


def extract_recommendation_table(table: List[List[str]], headers: List[str]) -> str:
    """Extract recommendation table with proper structure."""
    output_parts = []
    
    # Find key columns
    rec_col = action_col = owner_col = timescale_col = None
    
    for i, header in enumerate(headers):
        header_lower = header.lower()
        if 'recommendation' in header_lower and rec_col is None:
            rec_col = i
        elif 'action' in header_lower:
            action_col = i
        elif 'owner' in header_lower:
            owner_col = i
        elif 'timescale' in header_lower or 'time' in header_lower:
            timescale_col = i
    
    current_recommendation = None
    
    for row in table[1:]:
        col_indices = [c for c in [rec_col, action_col, owner_col, timescale_col] if c is not None]
        max_col = max(col_indices) if col_indices else 0
        if len(row) <= max_col:
            continue
        
        # Get recommendation text
        if rec_col is not None and len(row) > rec_col and row[rec_col].strip():
            if current_recommendation:
                output_parts.append(current_recommendation)
            
            rec_text = row[rec_col].strip()
            rec_match = re.match(r'^(\d+[a-z]?)\s*\.?\s*(.+)', rec_text)
            if rec_match:
                rec_num = rec_match.group(1)
                rec_content = rec_match.group(2)
                current_recommendation = f"Recommendation {rec_num}: {rec_content}"
            else:
                current_recommendation = f"Recommendation: {rec_text}"
        
        # Append actions
        if action_col is not None and len(row) > action_col and row[action_col].strip():
            action_text = row[action_col].strip()
            action_text = re.sub(r'^[a-z]\s+', '', action_text)
            if current_recommendation:
                current_recommendation += f" Action: {action_text}"
        
        # Add owner and timescale
        if owner_col is not None and len(row) > owner_col and row[owner_col].strip():
            if current_recommendation:
                current_recommendation += f" [Owner: {row[owner_col].strip()}]"
        
        if timescale_col is not None and len(row) > timescale_col and row[timescale_col].strip():
            if current_recommendation:
                current_recommendation += f" [Timescale: {row[timescale_col].strip()}]"
    
    if current_recommendation:
        output_parts.append(current_recommendation)
    
    return '\n\n'.join(output_parts)


def extract_generic_table(table: List[List[str]], headers: List[str]) -> str:
    """Extract generic table as structured text."""
    output_parts = []
    
    for row in table[1:]:
        row_parts = []
        for i, cell in enumerate(row):
            if cell.strip():
                header = headers[i] if i < len(headers) else f"Column {i+1}"
                if header.strip():
                    row_parts.append(f"{header}: {cell}")
                else:
                    row_parts.append(cell)
        
        if row_parts:
            output_parts.append(' | '.join(row_parts))
    
    return '\n'.join(output_parts)


def extract_pdf_text(uploaded_file) -> str:
    """
    ENHANCED: Extract text from PDF with better cleaning and deduplication
    """
    if not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
        raise ImportError("No PDF library available")
    
    try:
        if PDFPLUMBER_AVAILABLE:
            import pdfplumber
            with pdfplumber.open(uploaded_file) as pdf:
                text_parts = []
                seen_texts = set()  # NEW: Track seen text to avoid duplicates
                
                for page_num, page in enumerate(pdf.pages):
                    page_text_parts = []
                    
                    # Check for tables
                    tables = page.extract_tables()
                    has_substantial_tables = tables and any(
                        len(t) > 2 and len(t[0]) > 2 for t in tables if t
                    )
                    
                    if has_substantial_tables:
                        # Extract tables first
                        for table in tables:
                            if table and len(table) > 1:
                                table_text = extract_table_as_structured_text(table)
                                if table_text.strip():
                                    # NEW: Check for duplicates
                                    table_hash = hash(table_text.strip()[:100])
                                    if table_hash not in seen_texts:
                                        page_text_parts.append(table_text)
                                        seen_texts.add(table_hash)
                        
                        # Get non-table text
                        full_page_text = page.extract_text() or ""
                        non_table_lines = []
                        for line in full_page_text.split('\n'):
                            # Skip lines that look like table data
                            if not re.match(r'^(\S+\s{3,}){2,}', line):
                                non_table_lines.append(line)
                        
                        non_table_text = '\n'.join(non_table_lines)
                        if non_table_text.strip():
                            page_text_parts.insert(0, non_table_text)
                    else:
                        page_text = page.extract_text()
                        if page_text:
                            page_text_parts.append(page_text)
                    
                    if page_text_parts:
                        combined = '\n\n'.join(page_text_parts)
                        text_parts.append(combined)
                
                full_text = '\n\n'.join(text_parts)
                
                # NEW: Remove repeated sections (common in PDFs with headers/footers)
                full_text = remove_repeated_sections(full_text)
                
                # Clean up spacing
                full_text = re.sub(r' {2,}', ' ', full_text)
                full_text = re.sub(r'(?<!\n)\n(?!\n)', ' ', full_text)
                
                return clean_extracted_text(full_text)
        
        elif PYPDF2_AVAILABLE:
            import PyPDF2
            from io import BytesIO
            
            pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.getvalue()))
            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            full_text = '\n\n'.join(text_parts)
            full_text = re.sub(r' {2,}', ' ', full_text)
            full_text = re.sub(r'(?<!\n)\n(?!\n)', ' ', full_text)
            
            return clean_extracted_text(full_text)
    
    except Exception as e:
        logging.error(f"PDF extraction error: {e}")
        raise Exception(f"Failed to extract PDF text: {str(e)}")


def remove_repeated_sections(text: str, min_length: int = 50) -> str:
    """
    NEW: Remove sections that repeat multiple times (headers/footers/page numbers)
    """
    if not text or len(text) < 100:
        return text
    
    # Split into paragraphs
    paragraphs = text.split('\n\n')
    
    # Count paragraph occurrences
    para_counts = {}
    for para in paragraphs:
        para_clean = para.strip()
        if len(para_clean) >= min_length:
            para_counts[para_clean] = para_counts.get(para_clean, 0) + 1
    
    # Find paragraphs that appear more than once
    repeated = {para for para, count in para_counts.items() if count > 1}
    
    # Keep only the first occurrence of repeated paragraphs
    seen = set()
    filtered_paragraphs = []
    
    for para in paragraphs:
        para_clean = para.strip()
        
        if para_clean in repeated:
            if para_clean not in seen:
                filtered_paragraphs.append(para)
                seen.add(para_clean)
        else:
            filtered_paragraphs.append(para)
    
    return '\n\n'.join(filtered_paragraphs)


def extract_txt_text(uploaded_file) -> str:
    """Extract text from TXT file"""
    try:
        text = uploaded_file.getvalue().decode('utf-8')
    except UnicodeDecodeError:
        try:
            text = uploaded_file.getvalue().decode('latin-1')
        except UnicodeDecodeError:
            text = uploaded_file.getvalue().decode('utf-8', errors='ignore')
    return clean_extracted_text(text)


def extract_docx_text(uploaded_file) -> str:
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed")
    
    try:
        import docx
        from io import BytesIO
        
        doc = docx.Document(BytesIO(uploaded_file.getvalue()))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return clean_extracted_text('\n'.join(text_parts))
    
    except Exception as e:
        logging.error(f"DOCX extraction error: {e}")
        raise Exception(f"Failed to extract DOCX text: {str(e)}")


def get_processing_stats():
    """Get processing statistics"""
    return {
        'documents_processed': 0,
        'pdf_library': 'pdfplumber' if PDFPLUMBER_AVAILABLE else 'PyPDF2' if PYPDF2_AVAILABLE else 'none',
        'docx_library': 'python-docx' if DOCX_AVAILABLE else 'none',
        'table_extraction': 'enabled' if PDFPLUMBER_AVAILABLE else 'disabled'
    }


def check_dependencies():
    """Check available libraries"""
    return {
        'pdfplumber': PDFPLUMBER_AVAILABLE,
        'PyPDF2': PYPDF2_AVAILABLE,
        'python-docx': DOCX_AVAILABLE
    }
