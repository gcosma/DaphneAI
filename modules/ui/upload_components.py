# modules/ui/upload_components.py
# COMPLETE UPLOAD COMPONENTS - Fixes Document class import issues

import streamlit as st
import pandas as pd
import logging
import io
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

# Fix Document class import with comprehensive fallbacks
try:
    from langchain.schema import Document
    LANGCHAIN_AVAILABLE = True
except ImportError:
    try:
        from langchain_core.documents import Document
        LANGCHAIN_AVAILABLE = True
    except ImportError:
        # Create comprehensive Document class fallback
        class Document:
            def __init__(self, page_content: str, metadata: Dict = None):
                self.page_content = page_content
                self.metadata = metadata or {}
                
            def __str__(self):
                return self.page_content
            
            def __repr__(self):
                return f"Document(page_content='{self.page_content[:50]}...', metadata={self.metadata})"
        
        LANGCHAIN_AVAILABLE = False

# Import PDF processing libraries with fallbacks
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Import core utilities
try:
    from modules.core_utils import SecurityValidator, log_user_action
    CORE_UTILS_AVAILABLE = True
except ImportError:
    CORE_UTILS_AVAILABLE = False
    
    # Fallback security validator
    class SecurityValidator:
        def validate_file_upload(self, content: bytes, filename: str) -> Tuple[bool, str]:
            max_size = 200 * 1024 * 1024  # 200MB
            if len(content) > max_size:
                return False, f"File too large: {len(content)/1024/1024:.1f}MB"
            return True, "File validated"
    
    def log_user_action(action: str, details: Dict = None, user_id: str = "anonymous"):
        logging.info(f"USER_ACTION: {action} - {details}")

# ===============================================
# DOCUMENT PROCESSING CLASSES
# ===============================================

class DocumentProcessor:
    """Process uploaded documents with multiple format support"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.security_validator = SecurityValidator()
        self.supported_types = ['pdf', 'txt', 'docx', 'doc']
        
    def process_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
        """Process a single uploaded file"""
        try:
            # Read file content
            file_content = uploaded_file.read()
            filename = uploaded_file.name
            file_size = len(file_content)
            
            # Reset file pointer for potential re-reading
            uploaded_file.seek(0)
            
            # Validate file
            is_valid, validation_msg = self.security_validator.validate_file_upload(
                file_content, filename
            )
            
            if not is_valid:
                return {
                    'filename': filename,
                    'size': file_size,
                    'status': 'error',
                    'error': validation_msg,
                    'processed_at': datetime.now().isoformat()
                }
            
            # Extract text content
            extracted_text = self._extract_text_content(file_content, filename)
            
            # Create document object
            document = Document(
                page_content=extracted_text,
                metadata={
                    'filename': filename,
                    'file_size': file_size,
                    'upload_time': datetime.now().isoformat(),
                    'file_type': Path(filename).suffix.lower()
                }
            )
            
            return {
                'filename': filename,
                'size': file_size,
                'type': Path(filename).suffix.lower(),
                'content': extracted_text,
                'document_object': document,
                'status': 'success',
                'processed_at': datetime.now().isoformat(),
                'word_count': len(extracted_text.split()) if extracted_text else 0,
                'character_count': len(extracted_text) if extracted_text else 0
            }
            
        except Exception as e:
            self.logger.error(f"Error processing file {uploaded_file.name}: {e}")
            return {
                'filename': uploaded_file.name,
                'size': 0,
                'status': 'error',
                'error': str(e),
                'processed_at': datetime.now().isoformat()
            }
    
    def _extract_text_content(self, file_content: bytes, filename: str) -> str:
        """Extract text content from various file formats"""
        file_ext = Path(filename).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf_text(file_content)
            elif file_ext == '.txt':
                return file_content.decode('utf-8', errors='ignore')
            elif file_ext in ['.docx', '.doc']:
                return self._extract_word_text(file_content)
            else:
                # Try to decode as text
                return file_content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {filename}: {e}")
            return f"Error extracting text: {str(e)}"
    
    def _extract_pdf_text(self, pdf_content: bytes) -> str:
        """Extract text from PDF using available libraries"""
        text = ""
        
        # Try pdfplumber first (better for tables and structured content)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                
                if text.strip():
                    self.logger.info("Successfully extracted PDF text using pdfplumber")
                    return text.strip()
            except Exception as e:
                self.logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Fallback to PyMuPDF
        if PYMUPDF_AVAILABLE:
            try:
                pdf_doc = fitz.open(stream=pdf_content, filetype="pdf")
                for page_num in range(pdf_doc.page_count):
                    page = pdf_doc[page_num]
                    page_text = page.get_text()
                    if page_text:
                        text += page_text + "\n\n"
                pdf_doc.close()
                
                if text.strip():
                    self.logger.info("Successfully extracted PDF text using PyMuPDF")
                    return text.strip()
            except Exception as e:
                self.logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # If both libraries failed or are unavailable
        if not PDFPLUMBER_AVAILABLE and not PYMUPDF_AVAILABLE:
            return "PDF text extraction not available. Please install pdfplumber or PyMuPDF."
        
        return "Failed to extract text from PDF. The file may be corrupted or contain only images."
    
    def _extract_word_text(self, doc_content: bytes) -> str:
        """Extract text from Word documents"""
        try:
            # Try python-docx if available
            import docx
            doc = docx.Document(io.BytesIO(doc_content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except ImportError:
            return "Word document processing not available. Please install python-docx."
        except Exception as e:
            return f"Error extracting Word document text: {str(e)}"

# ===============================================
# UI COMPONENTS
# ===============================================

def render_upload_tab():
    """Main upload tab with comprehensive file handling"""
    st.header("📁 Document Upload")
    
    st.markdown("""
    Upload government inquiry reports, response documents, or related files for analysis. 
    Supported formats: PDF, TXT, DOCX, DOC
    """)
    
    # File uploader with proper configuration
    uploaded_files = st.file_uploader(
        "Choose files to upload",
        type=['pdf', 'txt', 'docx', 'doc'],
        accept_multiple_files=True,
        help="Upload up to 200MB per file (configurable in .streamlit/config.toml)",
        key="main_file_uploader"
    )
    
    if uploaded_files:
        render_file_processing_interface(uploaded_files)
    
    # Display previously uploaded documents
    render_uploaded_documents_summary()
    
    # Upload settings and options
    render_upload_settings()

def render_file_processing_interface(uploaded_files):
    """Process and display uploaded files"""
    st.subheader(f"📄 Processing {len(uploaded_files)} files...")
    
    # Initialize document processor
    processor = DocumentProcessor()
    
    # Progress tracking
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    processed_documents = []
    
    for i, uploaded_file in enumerate(uploaded_files):
        # Update progress
        progress = (i + 1) / len(uploaded_files)
        progress_bar.progress(progress)
        status_text.text(f"Processing {uploaded_file.name}...")
        
        # Process file
        processed_doc = processor.process_uploaded_file(uploaded_file)
        processed_documents.append(processed_doc)
        
        # Show individual file status
        if processed_doc['status'] == 'success':
            st.success(f"✅ {processed_doc['filename']} ({processed_doc['size']/1024/1024:.1f}MB)")
        else:
            st.error(f"❌ {processed_doc['filename']}: {processed_doc['error']}")
    
    # Update session state
    if 'uploaded_documents' not in st.session_state:
        st.session_state.uploaded_documents = []
    
    # Add new documents to session state
    successful_docs = [doc for doc in processed_documents if doc['status'] == 'success']
    st.session_state.uploaded_documents.extend(successful_docs)
    
    # Final status
    status_text.text("✅ Processing complete!")
    progress_bar.progress(1.0)
    
    # Display processing summary
    render_processing_summary(processed_documents)
    
    # Log upload action
    if CORE_UTILS_AVAILABLE:
        log_user_action(
            "documents_uploaded",
            {
                'file_count': len(uploaded_files),
                'successful_count': len(successful_docs),
                'total_size_mb': sum(doc['size'] for doc in successful_docs) / (1024 * 1024)
            }
        )

def render_processing_summary(processed_documents: List[Dict[str, Any]]):
    """Display summary of processed documents"""
    st.subheader("📊 Processing Summary")
    
    successful_docs = [doc for doc in processed_documents if doc['status'] == 'success']
    failed_docs = [doc for doc in processed_documents if doc['status'] == 'error']
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Files", len(processed_documents))
    
    with col2:
        st.metric("Successful", len(successful_docs))
    
    with col3:
        st.metric("Failed", len(failed_docs))
    
    with col4:
        total_size = sum(doc['size'] for doc in successful_docs) / (1024 * 1024)
        st.metric("Total Size", f"{total_size:.1f}MB")
    
    # Show failed files if any
    if failed_docs:
        st.warning("⚠️ Some files failed to process:")
        for doc in failed_docs:
            st.write(f"• **{doc['filename']}**: {doc['error']}")
    
    # Show successful files details
    if successful_docs:
        st.success("✅ Successfully processed files:")
        
        # Create summary table
        summary_data = []
        for doc in successful_docs:
            summary_data.append({
                'Filename': doc['filename'],
                'Size (MB)': f"{doc['size'] / (1024 * 1024):.1f}",
                'Type': doc['type'],
                'Words': doc.get('word_count', 'N/A'),
                'Characters': doc.get('character_count', 'N/A')
            })
        
        df = pd.DataFrame(summary_data)
        st.dataframe(df, use_container_width=True)

def render_uploaded_documents_summary():
    """Display summary of all uploaded documents"""
    if 'uploaded_documents' not in st.session_state or not st.session_state.uploaded_documents:
        st.info("📝 No documents uploaded yet.")
        return
    
    st.subheader("📚 Uploaded Documents")
    
    docs = st.session_state.uploaded_documents
    
    # Summary metrics
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Total Documents", len(docs))
    
    with col2:
        total_size = sum(doc.get('size', 0) for doc in docs) / (1024 * 1024)
        st.metric("Total Size", f"{total_size:.1f}MB")
    
    with col3:
        total_words = sum(doc.get('word_count', 0) for doc in docs)
        st.metric("Total Words", f"{total_words:,}")
    
    # Document list with actions
    for i, doc in enumerate(docs):
        with st.expander(f"📄 {doc['filename']} ({doc.get('size', 0)/1024/1024:.1f}MB)"):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                st.write(f"**Type:** {doc.get('type', 'Unknown')}")
                st.write(f"**Upload time:** {doc.get('processed_at', 'Unknown')}")
                st.write(f"**Words:** {doc.get('word_count', 'N/A'):,}")
                st.write(f"**Characters:** {doc.get('character_count', 'N/A'):,}")
                
                # Content preview
                content = doc.get('content', '')
                if content:
                    preview = content[:300] + "..." if len(content) > 300 else content
                    st.text_area("Content Preview", preview, height=100, disabled=True)
            
            with col2:
                if st.button(f"🗑️ Remove", key=f"remove_{i}"):
                    st.session_state.uploaded_documents.pop(i)
                    st.rerun()
                
                if st.button(f"⬇️ Download", key=f"download_{i}"):
                    st.download_button(
                        "Download Content",
                        doc.get('content', ''),
                        file_name=f"{Path(doc['filename']).stem}_extracted.txt",
                        mime="text/plain",
                        key=f"download_content_{i}"
                    )
    
    # Bulk actions
    if len(docs) > 1:
        st.subheader("🔧 Bulk Actions")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🗑️ Clear All Documents"):
                st.session_state.uploaded_documents = []
                st.rerun()
        
        with col2:
            if st.button("📥 Export All Content"):
                export_all_documents()
        
        with col3:
            if st.button("🔄 Refresh Analysis"):
                st.info("Documents refreshed! Ready for new analysis.")

def render_upload_settings():
    """Render upload configuration and settings"""
    with st.expander("⚙️ Upload Settings"):
        st.markdown("#### File Processing Options")
        
        col1, col2 = st.columns(2)
        
        with col1:
            auto_process = st.checkbox(
                "Auto-process on upload",
                value=True,
                help="Automatically extract text when files are uploaded"
            )
            
            preserve_formatting = st.checkbox(
                "Preserve formatting",
                value=False,
                help="Attempt to preserve document formatting (experimental)"
            )
        
        with col2:
            extract_metadata = st.checkbox(
                "Extract metadata",
                value=True,
                help="Extract file metadata and document properties"
            )
            
            validate_content = st.checkbox(
                "Validate content",
                value=True,
                help="Validate document content for analysis readiness"
            )
        
        # Advanced settings
        st.markdown("#### Advanced Settings")
        
        max_file_size = st.slider(
            "Max file size (MB)",
            min_value=1,
            max_value=500,
            value=200,
            help="Maximum allowed file size per document"
        )
        
        chunk_size = st.selectbox(
            "Processing chunk size",
            options=[1024, 2048, 4096, 8192],
            index=2,
            help="Size of chunks for processing large files"
        )
        
        # Store settings in session state
        st.session_state.upload_settings = {
            'auto_process': auto_process,
            'preserve_formatting': preserve_formatting,
            'extract_metadata': extract_metadata,
            'validate_content': validate_content,
            'max_file_size_mb': max_file_size,
            'chunk_size': chunk_size
        }

def render_document_validation():
    """Validate uploaded documents for analysis readiness"""
    if 'uploaded_documents' not in st.session_state or not st.session_state.uploaded_documents:
        return
    
    st.subheader("✅ Document Validation")
    
    docs = st.session_state.uploaded_documents
    validation_results = []
    
    for doc in docs:
        result = validate_document_for_analysis(doc)
        validation_results.append(result)
    
    # Summary of validation
    valid_count = sum(1 for r in validation_results if r['is_valid'])
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.metric("Valid Documents", f"{valid_count}/{len(docs)}")
    
    with col2:
        readiness_score = (valid_count / len(docs)) * 100 if docs else 0
        st.metric("Readiness Score", f"{readiness_score:.0f}%")
    
    # Detailed validation results
    for i, result in enumerate(validation_results):
        doc = docs[i]
        status_icon = "✅" if result['is_valid'] else "⚠️"
        
        with st.expander(f"{status_icon} {doc['filename']} - Validation Results"):
            
            if result['is_valid']:
                st.success("✅ Document is ready for analysis")
            else:
                st.warning("⚠️ Document has validation issues")
            
            # Show validation details
            for check, passed in result['checks'].items():
                check_icon = "✅" if passed else "❌"
                st.write(f"{check_icon} {check.replace('_', ' ').title()}")
            
            # Show recommendations if any issues
            if result['recommendations']:
                st.markdown("**Recommendations:**")
                for rec in result['recommendations']:
                    st.write(f"• {rec}")

def validate_document_for_analysis(doc: Dict[str, Any]) -> Dict[str, Any]:
    """Validate a document for analysis readiness"""
    checks = {}
    recommendations = []
    
    # Check content availability
    content = doc.get('content', '')
    checks['has_content'] = bool(content and content.strip())
    
    if not checks['has_content']:
        recommendations.append("Document appears to be empty or text extraction failed")
    
    # Check content length
    min_content_length = 100
    checks['sufficient_content'] = len(content) >= min_content_length
    
    if not checks['sufficient_content']:
        recommendations.append(f"Content is too short (less than {min_content_length} characters)")
    
    # Check for meaningful text (not just extraction errors)
    error_indicators = ['error extracting', 'failed to extract', 'not available']
    has_errors = any(indicator in content.lower() for indicator in error_indicators)
    checks['no_extraction_errors'] = not has_errors
    
    if has_errors:
        recommendations.append("Text extraction may have failed - check original file quality")
    
    # Check file type compatibility
    file_type = doc.get('type', '').lower()
    supported_types = ['.pdf', '.txt', '.docx', '.doc']
    checks['supported_format'] = file_type in supported_types
    
    if not checks['supported_format']:
        recommendations.append(f"File type {file_type} may not be fully supported")
    
    # Check for reasonable word count
    word_count = doc.get('word_count', 0)
    checks['reasonable_word_count'] = word_count >= 50
    
    if not checks['reasonable_word_count']:
        recommendations.append("Document has very few words - may not provide meaningful analysis")
    
    # Overall validation
    is_valid = all(checks.values())
    
    return {
        'is_valid': is_valid,
        'checks': checks,
        'recommendations': recommendations,
        'score': sum(checks.values()) / len(checks) if checks else 0
    }

def export_all_documents():
    """Export all uploaded documents content"""
    if 'uploaded_documents' not in st.session_state or not st.session_state.uploaded_documents:
        st.warning("No documents to export")
        return
    
    docs = st.session_state.uploaded_documents
    
    # Create combined content
    combined_content = []
    
    for doc in docs:
        combined_content.append(f"{'='*60}")
        combined_content.append(f"DOCUMENT: {doc['filename']}")
        combined_content.append(f"UPLOAD TIME: {doc.get('processed_at', 'Unknown')}")
        combined_content.append(f"SIZE: {doc.get('size', 0)/1024/1024:.1f}MB")
        combined_content.append(f"WORDS: {doc.get('word_count', 'N/A')}")
        combined_content.append(f"{'='*60}")
        combined_content.append("")
        combined_content.append(doc.get('content', 'No content available'))
        combined_content.append("")
        combined_content.append("")
    
    combined_text = "\n".join(combined_content)
    
    # Create download button
    st.download_button(
        "📥 Download All Documents",
        combined_text,
        file_name=f"all_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
        mime="text/plain"
    )
    
    # Also offer CSV export
    csv_data = []
    for doc in docs:
        csv_data.append({
            'Filename': doc['filename'],
            'Upload_Time': doc.get('processed_at', ''),
            'Size_MB': doc.get('size', 0) / (1024 * 1024),
            'Type': doc.get('type', ''),
            'Word_Count': doc.get('word_count', 0),
            'Character_Count': doc.get('character_count', 0),
            'Content': doc.get('content', '')[:1000] + "..." if len(doc.get('content', '')) > 1000 else doc.get('content', '')
        })
    
    if csv_data:
        df = pd.DataFrame(csv_data)
        csv_content = df.to_csv(index=False)
        
        st.download_button(
            "📊 Download Documents CSV",
            csv_content,
            file_name=f"documents_metadata_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )

def render_document_statistics():
    """Render comprehensive document statistics"""
    if 'uploaded_documents' not in st.session_state or not st.session_state.uploaded_documents:
        return
    
    st.subheader("📈 Document Statistics")
    
    docs = st.session_state.uploaded_documents
    
    # Calculate statistics
    stats = calculate_document_statistics(docs)
    
    # Display key metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Documents", stats['total_documents'])
    
    with col2:
        st.metric("Average Size", f"{stats['avg_size_mb']:.1f}MB")
    
    with col3:
        st.metric("Total Words", f"{stats['total_words']:,}")
    
    with col4:
        st.metric("Average Words", f"{stats['avg_words']:,.0f}")
    
    # File type distribution
    if stats['file_types']:
        st.markdown("#### 📊 File Type Distribution")
        
        type_df = pd.DataFrame([
            {'Type': file_type, 'Count': count}
            for file_type, count in stats['file_types'].items()
        ])
        
        st.bar_chart(type_df.set_index('Type'))
    
    # Document sizes
    if len(docs) > 1:
        st.markdown("#### 📏 Document Sizes")
        
        size_data = [doc.get('size', 0) / (1024 * 1024) for doc in docs]
        size_df = pd.DataFrame({
            'Document': [doc['filename'] for doc in docs],
            'Size (MB)': size_data
        })
        
        st.bar_chart(size_df.set_index('Document'))

def calculate_document_statistics(docs: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calculate comprehensive statistics for uploaded documents"""
    if not docs:
        return {}
    
    total_size = sum(doc.get('size', 0) for doc in docs)
    total_words = sum(doc.get('word_count', 0) for doc in docs)
    
    # File type distribution
    file_types = {}
    for doc in docs:
        file_type = doc.get('type', 'unknown')
        file_types[file_type] = file_types.get(file_type, 0) + 1
    
    # Size statistics
    sizes_mb = [doc.get('size', 0) / (1024 * 1024) for doc in docs]
    word_counts = [doc.get('word_count', 0) for doc in docs]
    
    return {
        'total_documents': len(docs),
        'total_size_mb': total_size / (1024 * 1024),
        'avg_size_mb': (total_size / (1024 * 1024)) / len(docs),
        'max_size_mb': max(sizes_mb) if sizes_mb else 0,
        'min_size_mb': min(sizes_mb) if sizes_mb else 0,
        'total_words': total_words,
        'avg_words': total_words / len(docs) if docs else 0,
        'max_words': max(word_counts) if word_counts else 0,
        'min_words': min(word_counts) if word_counts else 0,
        'file_types': file_types
    }

def render_upload_help():
    """Render help section for upload functionality"""
    with st.expander("❓ Upload Help & Tips"):
        st.markdown("""
        #### Supported File Formats
        - **PDF**: Government reports, inquiry documents, response documents
        - **TXT**: Plain text files with document content
        - **DOCX/DOC**: Microsoft Word documents
        
        #### Upload Tips
        - **File Size**: Up to 200MB per file (configurable)
        - **Quality**: Ensure PDFs contain text (not just scanned images)
        - **Multiple Files**: Upload multiple documents for batch processing
        - **File Names**: Use descriptive filenames for easier identification
        
        #### Troubleshooting
        - **Empty Content**: If PDF shows no content, it may be image-based
        - **Large Files**: Files over 200MB require configuration changes
        - **Processing Errors**: Check file integrity and try re-uploading
        
        #### Best Practices
        - Upload inquiry reports and government responses separately
        - Use consistent naming conventions for document types
        - Verify content extraction before proceeding to analysis
        """)

# ===============================================
# COMPONENT AVAILABILITY AND STATUS
# ===============================================

def get_upload_component_status() -> Dict[str, Any]:
    """Get status of upload component dependencies"""
    return {
        'langchain_available': LANGCHAIN_AVAILABLE,
        'pdfplumber_available': PDFPLUMBER_AVAILABLE,
        'pymupdf_available': PYMUPDF_AVAILABLE,
        'core_utils_available': CORE_UTILS_AVAILABLE,
        'pdf_processing_available': PDFPLUMBER_AVAILABLE or PYMUPDF_AVAILABLE
    }

def render_component_status():
    """Render component status for debugging"""
    if st.session_state.get('debug_mode', False):
        with st.expander("🔧 Upload Component Status"):
            status = get_upload_component_status()
            
            for component, available in status.items():
                icon = "✅" if available else "❌"
                st.write(f"{icon} {component.replace('_', ' ').title()}")
            
            if not status['pdf_processing_available']:
                st.warning("⚠️ No PDF processing libraries available. Install pdfplumber or PyMuPDF for PDF support.")

# ===============================================
# MAIN RENDER FUNCTION WITH ALL COMPONENTS
# ===============================================

def render_upload_tab():
    """Main upload tab with all functionality"""
    st.header("📁 Document Upload & Processing")
    
    # Component status (if debug mode)
    render_component_status()
    
    # File uploader
    uploaded_files = st.file_uploader(
        "Choose files to upload",
        type=['pdf', 'txt', 'docx', 'doc'],
        accept_multiple_files=True,
        help="Upload government documents for analysis (up to 200MB per file)",
        key="main_file_uploader"
    )
    
    # Process uploaded files
    if uploaded_files:
        render_file_processing_interface(uploaded_files)
    
    # Show uploaded documents summary
    if st.session_state.get('uploaded_documents'):
        render_uploaded_documents_summary()
        render_document_validation()
        render_document_statistics()
    
    # Settings and help
    col1, col2 = st.columns(2)
    
    with col1:
        render_upload_settings()
    
    with col2:
        render_upload_help()

# ===============================================
# EXPORTS
# ===============================================

__all__ = [
    'render_upload_tab',
    'DocumentProcessor',
    'render_file_processing_interface',
    'render_uploaded_documents_summary',
    'render_document_validation',
    'validate_document_for_analysis',
    'export_all_documents',
    'calculate_document_statistics',
    'get_upload_component_status',
    'Document',  # Export the Document class
    'LANGCHAIN_AVAILABLE',
    'PDFPLUMBER_AVAILABLE', 
    'PYMUPDF_AVAILABLE',
    'CORE_UTILS_AVAILABLE'
]
