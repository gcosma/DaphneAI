# app.py
# Universal AI-Powered Document Search Engine
# Works on ANY uploaded file with intelligent embeddings and semantic search

import streamlit as st
import sys
import time
import re
import numpy as np
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
import logging
import hashlib

# Setup
st.set_page_config(
    page_title="Universal AI Search",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# AI Dependencies with intelligent fallbacks
try:
    from sentence_transformers import SentenceTransformer
    import torch
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    EMBEDDINGS_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

# Document Processing - Universal support
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    from io import BytesIO
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Additional file type support
try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import json
    import xml.etree.ElementTree as ET
    STRUCTURED_AVAILABLE = True
except ImportError:
    STRUCTURED_AVAILABLE = False

class UniversalDocumentProcessor:
    """Universal document processor that handles ANY file type"""
    
    def __init__(self):
        self.supported_types = {
            'text': ['.txt', '.md', '.csv', '.log', '.json', '.xml', '.html', '.py', '.js', '.css'],
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc'],
            'excel': ['.xlsx', '.xls'],
            'other': []  # Will attempt to process as text
        }
    
    def process_file(self, uploaded_file) -> Dict[str, Any]:
        """Process any uploaded file and extract text content"""
        filename = uploaded_file.name
        file_ext = Path(filename).suffix.lower()
        file_size = len(uploaded_file.getvalue())
        
        try:
            # Determine processing method
            text_content = ""
            processing_method = "unknown"
            
            # PDF Processing (multiple fallbacks)
            if file_ext == '.pdf':
                text_content, processing_method = self._process_pdf(uploaded_file)
            
            # Microsoft Word Processing
            elif file_ext in ['.docx', '.doc'] and DOCX_AVAILABLE:
                text_content, processing_method = self._process_docx(uploaded_file)
            
            # Excel/Spreadsheet Processing
            elif file_ext in ['.xlsx', '.xls'] and EXCEL_AVAILABLE:
                text_content, processing_method = self._process_excel(uploaded_file)
            
            # Structured Data Processing
            elif file_ext in ['.json', '.xml']:
                text_content, processing_method = self._process_structured(uploaded_file, file_ext)
            
            # Plain Text Processing (universal fallback)
            else:
                text_content, processing_method = self._process_text(uploaded_file)
            
            if not text_content.strip():
                return {
                    'filename': filename,
                    'error': f'No extractable text content found in {file_ext} file',
                    'file_type': file_ext[1:] if file_ext else 'unknown',
                    'file_size_mb': file_size / 1024 / 1024
                }
            
            # Advanced text analysis
            words = text_content.split()
            sentences = len(re.findall(r'[.!?]+', text_content))
            paragraphs = len([p for p in text_content.split('\n\n') if p.strip()])
            
            # Generate content hash for deduplication
            content_hash = hashlib.md5(text_content.encode()).hexdigest()[:12]
            
            return {
                'filename': filename,
                'text': text_content,
                'content_hash': content_hash,
                'word_count': len(words),
                'sentence_count': sentences,
                'paragraph_count': paragraphs,
                'character_count': len(text_content),
                'file_type': file_ext[1:] if file_ext else 'unknown',
                'file_size_mb': file_size / 1024 / 1024,
                'processing_method': processing_method,
                'processed_at': datetime.now().isoformat(),
                'languages_detected': self._detect_languages(text_content),
                'content_preview': text_content[:500] + "..." if len(text_content) > 500 else text_content
            }
            
        except Exception as e:
            return {
                'filename': filename,
                'error': f'Processing failed: {str(e)}',
                'file_type': file_ext[1:] if file_ext else 'unknown',
                'file_size_mb': file_size / 1024 / 1024,
                'processing_method': 'failed'
            }
    
    def _process_pdf(self, uploaded_file) -> Tuple[str, str]:
        """Process PDF with multiple fallback methods"""
        # Method 1: pdfplumber (best quality)
        if PDF_AVAILABLE:
            try:
                with pdfplumber.open(uploaded_file) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    if text_parts:
                        return '\n\n'.join(text_parts), "pdfplumber"
            except Exception:
                pass
        
        # Method 2: PyPDF2 (fallback)
        if PYPDF2_AVAILABLE:
            try:
                from io import BytesIO
                uploaded_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.getvalue()))
                text_parts = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                if text_parts:
                    return '\n\n'.join(text_parts), "PyPDF2"
            except Exception:
                pass
        
        return "", "pdf_failed"
    
    def _process_docx(self, uploaded_file) -> Tuple[str, str]:
        """Process Microsoft Word documents"""
        try:
            doc = docx.Document(BytesIO(uploaded_file.getvalue()))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts), "python-docx"
        except Exception:
            return "", "docx_failed"
    
    def _process_excel(self, uploaded_file) -> Tuple[str, str]:
        """Process Excel/spreadsheet files"""
        try:
            # Try reading all sheets
            excel_file = pd.ExcelFile(uploaded_file)
            text_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
                
                # Convert dataframe to text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False, na_rep='')
                text_parts.append(sheet_text)
            
            return '\n\n'.join(text_parts), "pandas"
        except Exception:
            return "", "excel_failed"
    
    def _process_structured(self, uploaded_file, file_ext: str) -> Tuple[str, str]:
        """Process structured data files (JSON, XML)"""
        try:
            content = uploaded_file.getvalue().decode('utf-8')
            
            if file_ext == '.json':
                # Parse and format JSON
                data = json.loads(content)
                formatted = json.dumps(data, indent=2, ensure_ascii=False)
                return formatted, "json"
            
            elif file_ext == '.xml':
                # Parse and format XML
                root = ET.fromstring(content)
                text_parts = []
                
                def extract_xml_text(element, level=0):
                    indent = "  " * level
                    if element.text and element.text.strip():
                        text_parts.append(f"{indent}{element.tag}: {element.text.strip()}")
                    else:
                        text_parts.append(f"{indent}{element.tag}")
                    
                    for child in element:
                        extract_xml_text(child, level + 1)
                
                extract_xml_text(root)
                return '\n'.join(text_parts), "xml"
        
        except Exception:
            # Fallback to raw text
            try:
                return uploaded_file.getvalue().decode('utf-8'), f"{file_ext}_raw"
            except:
                return "", f"{file_ext}_failed"
    
    def _process_text(self, uploaded_file) -> Tuple[str, str]:
        """Process plain text files with encoding detection"""
        # Try multiple encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                content = uploaded_file.getvalue().decode(encoding)
                return content, f"text_{encoding}"
            except UnicodeDecodeError:
                continue
        
        # Last resort: decode with errors='ignore'
        try:
            content = uploaded_file.getvalue().decode('utf-8', errors='ignore')
            return content, "text_utf8_ignore"
        except:
            return "", "text_failed"
    
    def _detect_languages(self, text: str) -> List[str]:
        """Simple language detection"""
        # Basic language indicators
        languages = []
        
        if re.search(r'\b(the|and|or|but|in|on|at|to|for|of|with|by)\b', text.lower()):
            languages.append('english')
        
        if re.search(r'\b(le|la|les|et|ou|mais|dans|sur|à|pour|de|avec|par)\b', text.lower()):
            languages.append('french')
        
        if re.search(r'\b(der|die|das|und|oder|aber|in|auf|zu|für|von|mit|durch)\b', text.lower()):
            languages.append('german')
        
        return languages or ['unknown']

class UniversalAISearchEngine:
    """Universal AI-powered search engine using embeddings and intelligent algorithms"""
    
    def __init__(self):
        self.embedding_model = None
        self.embedding_cache = {}
        self.search_cache = {}
        
    def initialize_ai_models(self):
        """Initialize AI models for search"""
        if EMBEDDINGS_AVAILABLE and self.embedding_model is None:
            try:
                with st.spinner("🧠 Loading AI models for intelligent search..."):
                    self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                    st.success("✅ AI models loaded successfully")
                return True
            except Exception as e:
                st.error(f"Failed to load AI models: {e}")
                return False
        return EMBEDDINGS_AVAILABLE
    
    def create_document_embeddings(self, documents: List[Dict]) -> Dict:
        """Create embeddings for all documents with intelligent chunking"""
        if not self.initialize_ai_models():
            return {}
        
        embeddings_data = {
            'document_embeddings': [],
            'chunk_embeddings': [],
            'documents': documents,
            'chunks': [],
            'chunk_to_doc_map': []
        }
        
        try:
            with st.spinner("🔗 Creating AI embeddings for intelligent search..."):
                progress_bar = st.progress(0)
                
                all_chunks = []
                chunk_to_doc_map = []
                
                for doc_idx, doc in enumerate(documents):
                    if 'text' not in doc or not doc['text']:
                        continue
                    
                    # Create document-level embedding
                    doc_text = doc['text'][:2000]  # Limit for document embedding
                    doc_embedding = self.embedding_model.encode([doc_text])[0]
                    embeddings_data['document_embeddings'].append({
                        'embedding': doc_embedding,
                        'document_index': doc_idx,
                        'text_preview': doc_text[:200]
                    })
                    
                    # Create intelligent chunks
                    chunks = self._create_intelligent_chunks(doc['text'])
                    
                    for chunk in chunks:
                        all_chunks.append(chunk)
                        chunk_to_doc_map.append(doc_idx)
                    
                    progress_bar.progress((doc_idx + 1) / len(documents))
                
                # Create chunk embeddings
                if all_chunks:
                    chunk_embeddings = self.embedding_model.encode(all_chunks, show_progress_bar=True)
                    embeddings_data['chunk_embeddings'] = chunk_embeddings
                    embeddings_data['chunks'] = all_chunks
                    embeddings_data['chunk_to_doc_map'] = chunk_to_doc_map
                
                st.success(f"✅ Created embeddings for {len(embeddings_data['document_embeddings'])} documents and {len(all_chunks)} chunks")
                return embeddings_data
                
        except Exception as e:
            st.error(f"Failed to create embeddings: {e}")
            return {}
    
    def _create_intelligent_chunks(self, text: str, chunk_size: int = 400, overlap: int = 50) -> List[str]:
        """Create intelligent text chunks that preserve context"""
        if len(text.split()) <= chunk_size:
            return [text]
        
        # Try to split by sentences first
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        current_word_count = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            sentence_words = len(sentence.split())
            
            if current_word_count + sentence_words <= chunk_size:
                current_chunk += sentence + ". "
                current_word_count += sentence_words
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                if len(chunks) > 0 and overlap > 0:
                    # Add last few words from previous chunk
                    overlap_words = current_chunk.split()[-overlap:]
                    current_chunk = " ".join(overlap_words) + " " + sentence + ". "
                    current_word_count = len(overlap_words) + sentence_words
                else:
                    current_chunk = sentence + ". "
                    current_word_count = sentence_words
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def semantic_search(self, query: str, embeddings_data: Dict, max_results: int = 10) -> List[Dict]:
        """Advanced semantic search using AI embeddings"""
        if not embeddings_data or not self.embedding_model:
            return []
        
        try:
            # Create query embedding
            query_embedding = self.embedding_model.encode([query])[0]
            
            # Search at chunk level for detailed results
            chunk_results = []
            if embeddings_data.get('chunk_embeddings') is not None:
                chunk_embeddings = embeddings_data['chunk_embeddings']
                similarities = cosine_similarity([query_embedding], chunk_embeddings)[0]
                
                # Get top chunk matches
                top_chunk_indices = similarities.argsort()[-max_results * 3:][::-1]
                
                for idx in top_chunk_indices:
                    if similarities[idx] > 0.1:  # Minimum similarity threshold
                        doc_idx = embeddings_data['chunk_to_doc_map'][idx]
                        chunk_text = embeddings_data['chunks'][idx]
                        
                        chunk_results.append({
                            'document_index': doc_idx,
                            'similarity': similarities[idx],
                            'chunk_text': chunk_text,
                            'chunk_index': idx
                        })
            
            # Aggregate results by document
            doc_results = {}
            for result in chunk_results:
                doc_idx = result['document_index']
                if doc_idx not in doc_results:
                    doc_results[doc_idx] = {
                        'document': embeddings_data['documents'][doc_idx],
                        'max_similarity': result['similarity'],
                        'best_chunk': result['chunk_text'],
                        'chunk_count': 1,
                        'total_similarity': result['similarity']
                    }
                else:
                    doc_results[doc_idx]['chunk_count'] += 1
                    doc_results[doc_idx]['total_similarity'] += result['similarity']
                    if result['similarity'] > doc_results[doc_idx]['max_similarity']:
                        doc_results[doc_idx]['max_similarity'] = result['similarity']
                        doc_results[doc_idx]['best_chunk'] = result['chunk_text']
            
            # Format final results
            final_results = []
            for doc_data in doc_results.values():
                # Calculate composite score
                composite_score = (
                    doc_data['max_similarity'] * 0.7 +  # Best match weight
                    (doc_data['total_similarity'] / doc_data['chunk_count']) * 0.3  # Average weight
                )
                
                final_results.append({
                    'document': doc_data['document'],
                    'similarity': composite_score,
                    'max_chunk_similarity': doc_data['max_similarity'],
                    'chunk_matches': doc_data['chunk_count'],
                    'snippet': self._extract_semantic_snippet(doc_data['document']['text'], query, doc_data['best_chunk']),
                    'search_type': 'AI Semantic',
                    'confidence': self._calculate_confidence(composite_score),
                    'ai_explanation': self._generate_relevance_explanation(query, doc_data['best_chunk'], composite_score)
                })
            
            # Sort by composite score
            final_results.sort(key=lambda x: x['similarity'], reverse=True)
            return final_results[:max_results]
            
        except Exception as e:
            st.error(f"Semantic search failed: {e}")
            return []
    
    def hybrid_intelligent_search(self, query: str, documents: List[Dict], embeddings_data: Dict, max_results: int = 10) -> List[Dict]:
        """Combine AI semantic search with traditional methods for best results"""
        results = {}
        
        # 1. AI Semantic Search (primary method)
        if embeddings_data and self.embedding_model:
            semantic_results = self.semantic_search(query, embeddings_data, max_results)
            for result in semantic_results:
                doc_id = result['document']['filename']
                result['final_score'] = result['similarity'] * 1.2  # AI boost
                results[doc_id] = result
        
        # 2. Traditional keyword search (backup/enhancement)
        keyword_results = self.advanced_keyword_search(query, documents, max_results)
        for result in keyword_results:
            doc_id = result['document']['filename']
            if doc_id in results:
                # Enhance existing result
                results[doc_id]['final_score'] += result['score'] * 0.3
                results[doc_id]['search_type'] = 'AI + Keywords'
            else:
                # Add as keyword-only result
                result['final_score'] = result['score'] * 0.8
                result['search_type'] = 'Keywords Only'
                results[doc_id] = result
        
        # Sort by final score
        final_results = list(results.values())
        final_results.sort(key=lambda x: x.get('final_score', 0), reverse=True)
        
        return final_results[:max_results]
    
    def advanced_keyword_search(self, query: str, documents: List[Dict], max_results: int = 10) -> List[Dict]:
        """Advanced keyword search with intelligent scoring"""
        results = []
        query_lower = query.lower()
        query_words = query_lower.split()
        
        for doc in documents:
            if 'text' not in doc or not doc['text']:
                continue
            
            text = doc['text']
            text_lower = text.lower()
            
            # Multi-factor scoring
            score = 0
            
            # Exact phrase matching
            if query_lower in text_lower:
                score += text_lower.count(query_lower) * 10
            
            # Word matching with proximity bonus
            word_matches = 0
            for word in query_words:
                if word in text_lower:
                    word_matches += 1
                    # Frequency bonus
                    score += text_lower.count(word) * 2
            
            # Word coverage bonus
            if query_words:
                coverage = word_matches / len(query_words)
                score += coverage * 5
            
            # Position bonus (earlier = better)
            first_occurrence = text_lower.find(query_lower)
            if first_occurrence != -1:
                position_bonus = max(0, 2 - (first_occurrence / len(text_lower)))
                score += position_bonus
            
            if score > 0:
                snippet = self._extract_keyword_snippet(text, query)
                
                results.append({
                    'document': doc,
                    'score': score,
                    'snippet': snippet,
                    'search_type': 'Advanced Keywords',
                    'confidence': self._calculate_confidence(score / 10)  # Normalize for confidence
                })
        
        results.sort(key=lambda x: x['score'], reverse=True)
        return results[:max_results]
    
    def _extract_semantic_snippet(self, text: str, query: str, best_chunk: str) -> str:
        """Extract snippet for semantic search results"""
        # Find the best chunk in the full text
        chunk_start = text.find(best_chunk[:100])  # Use first 100 chars to locate
        
        if chunk_start != -1:
            # Extract with context
            start = max(0, chunk_start - 50)
            end = min(len(text), chunk_start + len(best_chunk) + 50)
            snippet = text[start:end]
        else:
            # Fallback to best chunk
            snippet = best_chunk[:400]
        
        return ("..." if chunk_start > 50 else "") + snippet + ("..." if len(snippet) < len(text) else "")
    
    def _extract_keyword_snippet(self, text: str, query: str, max_length: int = 300) -> str:
        """Extract snippet for keyword search"""
        query_lower = query.lower()
        text_lower = text.lower()
        
        # Find best position
        index = text_lower.find(query_lower)
        if index == -1:
            # Find first word
            for word in query_lower.split():
                index = text_lower.find(word)
                if index != -1:
                    break
        
        if index == -1:
            return text[:max_length] + "..."
        
        # Extract around found position
        start = max(0, index - max_length // 2)
        end = min(len(text), index + len(query) + max_length // 2)
        snippet = text[start:end]
        
        # Highlight query terms
        for word in query.split():
            snippet = re.sub(f'({re.escape(word)})', r'**\1**', snippet, flags=re.IGNORECASE)
        
        return ("..." if start > 0 else "") + snippet + ("..." if end < len(text) else "")
    
    def _calculate_confidence(self, score: float) -> str:
        """Calculate search confidence level"""
        if score >= 0.7:
            return "High"
        elif score >= 0.4:
            return "Medium"
        else:
            return "Low"
    
    def _generate_relevance_explanation(self, query: str, best_chunk: str, score: float) -> str:
        """Generate AI explanation of why result is relevant"""
        explanations = []
        
        query_words = query.lower().split()
        chunk_lower = best_chunk.lower()
        
        # Check for exact matches
        if query.lower() in chunk_lower:
            explanations.append("Contains exact query phrase")
        
        # Check word matches
        matched_words = [word for word in query_words if word in chunk_lower]
        if matched_words:
            explanations.append(f"Matches {len(matched_words)}/{len(query_words)} query terms")
        
        # Confidence level
        if score > 0.7:
            explanations.append("High semantic similarity")
        elif score > 0.4:
            explanations.append("Good semantic similarity")
        
        return " • ".join(explanations) if explanations else "Semantic content similarity"

def initialize_session_state():
    """Initialize comprehensive session state"""
    if 'initialized' not in st.session_state:
        st.session_state.initialized = True
        st.session_state.documents = []
        st.session_state.embeddings_data = {}
        st.session_state.search_engine = UniversalAISearchEngine()
        st.session_state.processor = UniversalDocumentProcessor()
        st.session_state.search_history = []
        st.session_state.analytics = {
            'total_searches': 0,
            'files_processed': 0,
            'ai_searches': 0,
            'keyword_searches': 0
        }

def render_universal_upload():
    """Universal file upload interface"""
    st.header("📁 Universal Document Upload")
    st.markdown("**Upload ANY file type - the AI will intelligently extract and process content**")
    
    # File type support info
    with st.expander("📋 Supported File Types", expanded=False):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Fully Supported:**")
            st.markdown("• 📄 PDF files")
            st.markdown("• 📝 Text files (.txt, .md, .log)")
            st.markdown("• 📊 Microsoft Word (.docx)")
            st.markdown("• 📈 Excel files (.xlsx, .xls)")
            st.markdown("• 🔗 JSON, XML files")
            
        with col2:
            st.markdown("**Experimental Support:**")
            st.markdown("• 💻 Code files (.py, .js, .html)")
            st.markdown("• 📋 CSV files")
            st.markdown("• 📰 HTML files")
            st.markdown("• 🔧 Configuration files")
            st.markdown("• 📄 Any text-based format")
    
    uploaded_files = st.file_uploader(
        "Drop files here or click to browse",
        accept_multiple_files=True,
        help="Upload any file type - AI will automatically detect and process content"
    )
    
    if uploaded_files:
        col1, col2 = st.columns(2)
        
        with col1:
            process_mode = st.radio(
                "Processing Mode:",
                ["🚀 Fast Processing", "🧠 AI-Enhanced Processing"],
                help="Fast: Quick text extraction | AI-Enhanced: Full AI analysis"
            )
        
        with col2:
            if st.button("🔄 Process All Files", type="primary"):
                process_universal_files(uploaded_files, process_mode)

def process_universal_files(uploaded_files, process_mode: str):
    """Process uploaded files with universal compatibility"""
    processor = st.session_state.processor
    processed_docs = []
    
    # Processing progress
    progress_bar = st.progress(0)
    status_container = st.container()
    
    for i, file in enumerate(uploaded_files):
        with status_container:
            st.text(f"🔄 Processing: {file.name}")
        
        # Process file
        doc_result = processor.process_file(file)
        processed_docs.append(doc_result)
        
        progress_bar.progress((i + 1) / len(uploaded_files))
    
    # Update session state
    st.session_state.documents = processed_docs
    st.session_state.embeddings_data = {}  # Reset embeddings
    st.session_state.analytics['files_processed'] += len(uploaded_files)
    
    # Show results
    with status_container:
        st.success(f"✅ Processed {len(uploaded_files)} files!")
        
        # Results summary
        successful = [doc for doc in processed_docs if 'error' not in doc]
        failed = [doc for doc in processed_docs if 'error' in doc]
        
        if successful:
            st.success(f"✅ Successfully processed: {len(successful)} files")
            
            # Show processing details
            with st.expander("📊 Processing Details", expanded=True):
                for doc in successful:
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.write(f"📄 **{doc['filename']}**")
                    with col2:
                        st.write(f"📝 {doc['word_count']:,} words")
                    with col3:
                        st.write(f"💾 {doc['file_size_mb']:.1f} MB")
                    with col4:
                        st.write(f"⚙️ {doc['processing_method']}")
        
        if failed:
            st.error(f"❌ Failed to process: {len(failed)} files")
            for doc in failed:
                st.error(f"• {doc['filename']}: {doc['error']}")
    
    # AI Enhancement option
    if process_mode == "🧠 AI-Enhanced Processing" and successful:
        if st.button("🧠 Build AI Search Index"):
            build_ai_search_index(successful)

def build_ai_search_index(documents: List[Dict]):
    """Build AI search index for enhanced search capabilities"""
    search_engine = st.session_state.search_engine
    
    if not EMBEDDINGS_AVAILABLE:
        st.warning("🤖 AI search requires additional packages. Install: pip install sentence-transformers torch scikit-learn")
        return
    
    with st.spinner("🧠 Building AI search index..."):
        embeddings_data = search_engine.create_document_embeddings(documents)
        st.session_state.embeddings_data = embeddings_data
        
        if embeddings_data:
            st.success("🎯 AI search index built successfully!")
        else:
            st.error("❌ Failed to build AI search index")

def render_universal_search():
    """Universal AI-powered search interface"""
    st.header("🔍 Universal AI Search")
    
    if not st.session_state.documents:
        st.info("📁 Please upload documents first")
        return
    
    # Document status
    valid_docs = [doc for doc in st.session_state.documents if 'text' in doc]
    st.info(f"📚 Ready to search {len(valid_docs)} documents")
    
    search_engine = st.session_state.search_engine
    
    # Search input
    query = st.text_input(
        "🔍 Search Query",
        placeholder="Ask anything about your documents...",
        help="Use natural language - AI understands context and meaning"
    )
    
    # Search configuration
    col1, col2, col3 = st.columns(3)
    
    with col1:
        search_modes = {
            'ai_hybrid': '🧠 AI Hybrid (Recommended)',
            'ai_semantic': '🤖 AI Semantic Only',
            'keywords': '🔍 Advanced Keywords',
            'auto': '⚡ Auto-Select Best'
        }
        
        search_mode = st.selectbox(
            "Search Method",
            options=list(search_modes.keys()),
            format_func=lambda x: search_modes[x],
            help="AI methods understand meaning, keywords find exact matches"
        )
    
    with col2:
        max_results = st.slider("Max Results", 1, 20, 8)
    
    with col3:
        show_explanations = st.checkbox("Show AI Explanations", value=True)
    
    # Advanced options
    with st.expander("⚙️ Advanced Search Options", expanded=False):
        col1, col2, col3 = st.columns(3)
        
        with col1:
            file_type_filter = st.multiselect(
                "File Types",
                options=list(set(doc.get('file_type', 'unknown') for doc in valid_docs)),
                help="Filter by file type"
            )
        
        with col2:
            min_similarity = st.slider("Min Similarity", 0.0, 1.0, 0.1, 0.1)
        
        with col3:
            content_filter = st.text_input("Content Filter", placeholder="Must contain...")
    
    # Perform search
    if query:
        # Apply filters
        filtered_docs = valid_docs
        if file_type_filter:
            filtered_docs = [doc for doc in filtered_docs if doc.get('file_type') in file_type_filter]
        if content_filter:
            filtered_docs = [doc for doc in filtered_docs if content_filter.lower() in doc.get('text', '').lower()]
        
        if not filtered_docs:
            st.warning("No documents match your filters")
            return
        
        # Execute search
        start_time = time.time()
        
        try:
            results = execute_universal_search(
                query, search_mode, filtered_docs, max_results, min_similarity
            )
            
            search_time = time.time() - start_time
            
            # Update analytics
            st.session_state.analytics['total_searches'] += 1
            if 'ai' in search_mode:
                st.session_state.analytics['ai_searches'] += 1
            else:
                st.session_state.analytics['keyword_searches'] += 1
            
            # Add to search history
            st.session_state.search_history.insert(0, {
                'query': query,
                'mode': search_mode,
                'results_count': len(results),
                'search_time': search_time,
                'timestamp': datetime.now().isoformat()
            })
            st.session_state.search_history = st.session_state.search_history[:10]
            
            # Display results
            display_universal_results(results, query, search_time, show_explanations)
            
        except Exception as e:
            st.error(f"Search failed: {e}")
            st.write("**Debug info:**", str(e))

def execute_universal_search(query: str, mode: str, documents: List[Dict], max_results: int, min_similarity: float) -> List[Dict]:
    """Execute search based on selected mode"""
    search_engine = st.session_state.search_engine
    embeddings_data = st.session_state.embeddings_data
    
    if mode == 'auto':
        # Auto-select best method
        if embeddings_data and EMBEDDINGS_AVAILABLE:
            mode = 'ai_hybrid'
        else:
            mode = 'keywords'
    
    if mode == 'ai_hybrid':
        if not embeddings_data:
            st.warning("🧠 Building AI index for first search...")
            embeddings_data = search_engine.create_document_embeddings(documents)
            st.session_state.embeddings_data = embeddings_data
        
        if embeddings_data:
            results = search_engine.hybrid_intelligent_search(query, documents, embeddings_data, max_results)
        else:
            results = search_engine.advanced_keyword_search(query, documents, max_results)
    
    elif mode == 'ai_semantic':
        if not embeddings_data:
            st.warning("🧠 Building AI index for semantic search...")
            embeddings_data = search_engine.create_document_embeddings(documents)
            st.session_state.embeddings_data = embeddings_data
        
        if embeddings_data:
            results = search_engine.semantic_search(query, embeddings_data, max_results)
        else:
            st.error("AI semantic search unavailable")
            results = []
    
    else:  # keywords
        results = search_engine.advanced_keyword_search(query, documents, max_results)
    
    # Apply similarity filter
    if min_similarity > 0:
        results = [r for r in results if r.get('similarity', r.get('score', 0)) >= min_similarity]
    
    return results

def display_universal_results(results: List[Dict], query: str, search_time: float, show_explanations: bool):
    """Display search results with rich information"""
    
    if not results:
        st.warning(f"❌ No results found for '{query}'")
        
        # Search suggestions
        st.markdown("### 💡 Search Tips:")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**For AI Search:**")
            st.markdown("• Use natural language")
            st.markdown("• Ask questions about content")
            st.markdown("• Describe concepts you're looking for")
        
        with col2:
            st.markdown("**For Better Results:**")
            st.markdown("• Try different keywords")
            st.markdown("• Use broader search terms")
            st.markdown("• Check file type filters")
        
        return
    
    # Results header
    st.success(f"🎯 Found {len(results)} results in {search_time:.2f} seconds")
    
    # Quick stats
    if len(results) > 1:
        with st.expander("📊 Search Statistics", expanded=False):
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                avg_score = np.mean([r.get('similarity', r.get('score', 0)) for r in results])
                st.metric("Avg Relevance", f"{avg_score:.2f}")
            
            with col2:
                file_types = [r['document'].get('file_type', 'unknown') for r in results]
                unique_types = len(set(file_types))
                st.metric("File Types", unique_types)
            
            with col3:
                ai_results = sum(1 for r in results if 'AI' in r.get('search_type', ''))
                st.metric("AI Results", ai_results)
            
            with col4:
                high_conf = sum(1 for r in results if r.get('confidence') == 'High')
                st.metric("High Confidence", high_conf)
    
    # Display each result
    for i, result in enumerate(results, 1):
        doc = result['document']
        
        with st.container():
            # Result header with enhanced info
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col1:
                st.markdown(f"### 📄 {i}. {doc['filename']}")
                
                # File info chips
                chips = []
                chips.append(f"📁 {doc.get('file_type', 'unknown').upper()}")
                chips.append(f"📝 {doc.get('word_count', 0):,} words")
                chips.append(f"💾 {doc.get('file_size_mb', 0):.1f} MB")
                if doc.get('processing_method'):
                    chips.append(f"⚙️ {doc['processing_method']}")
                
                st.caption(" • ".join(chips))
            
            with col2:
                # Confidence indicator
                confidence = result.get('confidence', 'Medium')
                confidence_colors = {'High': '🟢', 'Medium': '🟡', 'Low': '🔴'}
                st.markdown(f"**{confidence_colors.get(confidence, '⚪')} {confidence}**")
                
                # Search method
                search_type = result.get('search_type', 'Unknown')
                st.caption(f"Method: {search_type}")
            
            with col3:
                # Score display
                score = result.get('similarity', result.get('final_score', result.get('score', 0)))
                st.markdown(f"**Score: {score:.3f}**")
                
                # Additional metrics for AI results
                if 'chunk_matches' in result:
                    st.caption(f"Chunks: {result['chunk_matches']}")
            
            # AI Explanation
            if show_explanations and 'ai_explanation' in result:
                st.info(f"🧠 **Why relevant:** {result['ai_explanation']}")
            
            # Content preview
            snippet = result.get('snippet', '')
            if snippet:
                st.markdown("**📋 Relevant excerpt:**")
                st.markdown(snippet)
            elif 'content_preview' in doc:
                st.markdown("**📋 Content preview:**")
                st.markdown(doc['content_preview'])
            
            # Actions
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button(f"📖 Full Content", key=f"view_{i}"):
                    with st.expander(f"📄 Full content: {doc['filename']}", expanded=True):
                        st.text_area(
                            "Full Text",
                            doc.get('text', 'No content available'),
                            height=400,
                            key=f"content_{i}"
                        )
            
            with col2:
                if st.button(f"🔍 Find Similar", key=f"similar_{i}"):
                    find_similar_content(doc, query)
            
            with col3:
                if st.button(f"📊 Analysis", key=f"analyze_{i}"):
                    show_document_analysis(doc)
            
            st.markdown("---")

def find_similar_content(reference_doc: Dict, original_query: str):
    """Find documents similar to the selected one"""
    st.info("🔍 Finding similar content...")
    
    search_engine = st.session_state.search_engine
    
    # Use document content as query for similarity search
    doc_text = reference_doc.get('text', '')[:500]  # Use first 500 chars as query
    
    # Search for similar documents
    if st.session_state.embeddings_data:
        similar_results = search_engine.semantic_search(
            doc_text, 
            st.session_state.embeddings_data, 
            max_results=5
        )
        
        # Filter out the reference document itself
        similar_results = [r for r in similar_results if r['document']['filename'] != reference_doc['filename']]
        
        if similar_results:
            st.success(f"Found {len(similar_results)} similar documents:")
            for i, result in enumerate(similar_results, 1):
                st.write(f"{i}. **{result['document']['filename']}** (Similarity: {result['similarity']:.3f})")
        else:
            st.info("No similar documents found")
    else:
        st.warning("AI similarity search requires building the search index first")

def show_document_analysis(doc: Dict):
    """Show detailed document analysis"""
    st.info("📊 Document Analysis")
    
    analysis_data = {
        "Basic Stats": {
            "Word Count": f"{doc.get('word_count', 0):,}",
            "Character Count": f"{doc.get('character_count', 0):,}",
            "Sentences": doc.get('sentence_count', 0),
            "Paragraphs": doc.get('paragraph_count', 0)
        },
        "File Info": {
            "File Type": doc.get('file_type', 'unknown').upper(),
            "File Size": f"{doc.get('file_size_mb', 0):.2f} MB",
            "Processing Method": doc.get('processing_method', 'unknown'),
            "Processed At": doc.get('processed_at', 'unknown')
        }
    }
    
    # Display analysis
    for category, stats in analysis_data.items():
        st.write(f"**{category}:**")
        for key, value in stats.items():
            st.write(f"• {key}: {value}")

def render_analytics_dashboard():
    """Render comprehensive analytics dashboard"""
    st.header("📈 Analytics Dashboard")
    
    analytics = st.session_state.analytics
    
    # Main metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Searches", analytics['total_searches'])
    
    with col2:
        st.metric("Files Processed", analytics['files_processed'])
    
    with col3:
        st.metric("AI Searches", analytics['ai_searches'])
    
    with col4:
        st.metric("Keyword Searches", analytics['keyword_searches'])
    
    # Document statistics
    if st.session_state.documents:
        st.subheader("📚 Document Statistics")
        
        valid_docs = [doc for doc in st.session_state.documents if 'text' in doc]
        
        col1, col2 = st.columns(2)
        
        with col1:
            # File type distribution
            file_types = [doc.get('file_type', 'unknown') for doc in valid_docs]
            type_counts = {}
            for ft in file_types:
                type_counts[ft] = type_counts.get(ft, 0) + 1
            
            st.write("**File Types:**")
            for file_type, count in sorted(type_counts.items()):
                percentage = (count / len(valid_docs)) * 100
                st.write(f"• {file_type.upper()}: {count} ({percentage:.1f}%)")
        
        with col2:
            # Size distribution
            total_size = sum(doc.get('file_size_mb', 0) for doc in valid_docs)
            total_words = sum(doc.get('word_count', 0) for doc in valid_docs)
            
            st.write("**Content Statistics:**")
            st.write(f"• Total Size: {total_size:.2f} MB")
            st.write(f"• Total Words: {total_words:,}")
            st.write(f"• Avg Words/Doc: {total_words/len(valid_docs):,.0f}")
            st.write(f"• Avg Size/Doc: {total_size/len(valid_docs):.2f} MB")
    
    # Search history
    if st.session_state.search_history:
        st.subheader("🕒 Recent Searches")
        
        for i, search in enumerate(st.session_state.search_history[:5]):
            with st.expander(f"Search {i+1}: '{search['query']}'"):
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.write(f"**Method:** {search['mode']}")
                with col2:
                    st.write(f"**Results:** {search['results_count']}")
                with col3:
                    st.write(f"**Time:** {search['search_time']:.2f}s")
                
                if st.button(f"🔄 Repeat Search", key=f"repeat_{i}"):
                    st.session_state.repeat_query = search['query']
                    st.rerun()
    
    # System status
    st.subheader("🔧 System Status")
    
    status_items = [
        ("AI Embeddings", EMBEDDINGS_AVAILABLE, "sentence-transformers torch"),
        ("Advanced Analytics", SKLEARN_AVAILABLE, "scikit-learn"),
        ("PDF Processing", PDF_AVAILABLE, "pdfplumber"),
        ("Word Documents", DOCX_AVAILABLE, "python-docx"),
        ("Excel Files", EXCEL_AVAILABLE, "pandas openpyxl"),
        ("PyPDF2 Fallback", PYPDF2_AVAILABLE, "PyPDF2")
    ]
    
    for name, available, package in status_items:
        col1, col2 = st.columns([3, 1])
        with col1:
            if available:
                st.success(f"✅ {name}")
            else:
                st.error(f"❌ {name}")
                st.caption(f"Install: pip install {package}")
        with col2:
            st.write("Available" if available else "Missing")

def main():
    """Main application with universal AI search"""
    initialize_session_state()
    
    st.title("🧠 Universal AI Document Search")
    st.markdown("""
    **Intelligent search for ANY file type using advanced AI and embeddings**
    
    Upload any document and search with natural language understanding.
    """)
    
    # Sidebar with system status
    with st.sidebar:
        st.markdown("### 🔧 AI Capabilities")
        
        if EMBEDDINGS_AVAILABLE:
            st.success("🧠 AI Semantic Search")
        else:
            st.error("❌ AI Search Unavailable")
            st.caption("pip install sentence-transformers torch")
        
        if SKLEARN_AVAILABLE:
            st.success("📊 Advanced Analytics")
        else:
            st.warning("⚠️ Limited Analytics")
        
        # Quick stats
        if st.session_state.documents:
            st.markdown("### 📊 Quick Stats")
            valid_docs = [doc for doc in st.session_state.documents if 'text' in doc]
            st.metric("Documents", len(valid_docs))
            
            if st.session_state.embeddings_data:
                st.success("🔗 AI Index Built")
            else:
                st.info("🔗 AI Index Not Built")
            
            total_words = sum(doc.get('word_count', 0) for doc in valid_docs)
            st.metric("Total Words", f"{total_words:,}")
        
        # Handle repeat search
        if hasattr(st.session_state, 'repeat_query'):
            st.info(f"🔄 Repeating search: '{st.session_state.repeat_query}'")
            del st.session_state.repeat_query
    
    # Main interface tabs
    tab1, tab2, tab3 = st.tabs(["📁 Universal Upload", "🔍 AI Search", "📈 Analytics"])
    
    with tab1:
        render_universal_upload()
    
    with tab2:
        render_universal_search()
    
    with tab3:
        render_analytics_dashboard()

if __name__ == "__main__":
    main()
